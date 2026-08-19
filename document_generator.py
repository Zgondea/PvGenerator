import os
import io
import zipfile
import docx
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm

TEMPLATES_DIR = "templates"

def get_available_templates():
    """Recuperează lista de șabloane .docx disponibile în directorul templates."""
    if not os.path.exists(TEMPLATES_DIR):
        os.makedirs(TEMPLATES_DIR)
    
    templates = {}
    for file in os.listdir(TEMPLATES_DIR):
        if file.endswith(".docx") and not file.startswith("~$"):
            name = os.path.splitext(file)[0].replace("_", " ").title()
            templates[name] = os.path.join(TEMPLATES_DIR, file)
    return templates

def genereaza_documente(sabloane_selectate, dict_templates, context_baza, cale_sigla, inaltime_sigla=16):
    """Generează fișierele Word completate pe baza contextului."""
    rezultate = []

    for nume_sablon in sabloane_selectate:
        if nume_sablon in dict_templates:
            cale_template = dict_templates[nume_sablon]
            doc = DocxTemplate(cale_template)

            # Copie locală context pentru a adăuga imaginea fără a afecta celelalte șabloane
            context = context_baza.copy()

            if cale_sigla and os.path.exists(cale_sigla):
                context['sigla'] = InlineImage(doc, cale_sigla, height=Mm(inaltime_sigla))
            else:
                context['sigla'] = ""

            doc.render(context)

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            nume_fisier_iesire = f"{nume_sablon}.docx"
            rezultate.append((nume_fisier_iesire, buffer))

    return rezultate

def creeaza_arhiva_zip(fisiere_generate):
    """Impachetează multiple fișiere generate într-o arhivă ZIP în memorie."""
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for filename, buffer in fisiere_generate:
            zip_file.writestr(filename, buffer.getvalue())
    zip_buffer.seek(0)
    return zip_buffer

def render_docx_preview_html(docx_buffer):
    """
    Citește buffer-ul Word (.docx) și îl convertește într-o structură
    HTML curată, cu stiluri de document A4 pentru Live Preview.
    """
    docx_buffer.seek(0)
    doc = docx.Document(docx_buffer)
    
    html_out = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            style_name = paragraph.style.name.lower()
            if 'heading 1' in style_name or 'title' in style_name:
                html_out.append(f"<h2 style='color: #111111; border-bottom: 2px solid #333333; margin-top: 15px;'>{paragraph.text}</h2>")
            elif 'heading 2' in style_name:
                html_out.append(f"<h3 style='color: #222222; margin-top: 10px;'>{paragraph.text}</h3>")
            else:
                html_out.append(f"<p style='margin-bottom: 8px; color: #333333; line-height: 1.5;'>{paragraph.text}</p>")

    for table in doc.tables:
        table_html = "<table style='width:100%; border-collapse: collapse; margin: 15px 0; font-size: 12px; color: #111111;'>"
        for row in table.rows:
            table_html += "<tr>"
            for cell in row.cells:
                table_html += f"<td style='border: 1px solid #cccccc; padding: 6px;'>{cell.text}</td>"
            table_html += "</tr>"
        table_html += "</table>"
        html_out.append(table_html)

    return "".join(html_out)